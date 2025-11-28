import streamlit as st
import os
import docx
import requests
import time
from typing import TypedDict
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate
from langgraph.graph import StateGraph, END

# ==========================================
# 설정 및 비밀키 로드
# ==========================================
st.set_page_config(page_title="AI 채용공고 이력서 매칭 서비스", page_icon="🐈‍⬛", layout="wide")

if "GOOGLE_API_KEY" in st.secrets:
    os.environ["GOOGLE_API_KEY"] = st.secrets["GOOGLE_API_KEY"]
else:
    st.error("🚨 API 키가 없습니다. 개발자에게 문의 부탁드립니다")
    st.stop()

# 모델 : gemini-2.5-flash
llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0)

# ==========================================
# 디스코드 메세지 (시간 분석 / 별점)
# ==========================================
def send_discord_alert(user_name, duties, result_text, latency_ms):
    webhook_url = st.secrets.get("DISCORD_WEBHOOK_URL", None)
    if not webhook_url: return

    short_duties = duties[:200] + "..." if len(duties) > 200 else duties
    short_result = result_text[:800] + "\n..." if len(result_text) > 800 else result_text

    payload = {
        "username": "AI 매칭 분석 로그",
        "embeds": [{
            "title": "🚀 새로운 이용자 등장!",
            "color": 3447003, # 파란색
            "fields": [
                {"name": "사용자", "value": user_name if user_name else "익명", "inline": True},
                {"name": "처리 시간 (ms)", "value": f"{latency_ms:,} ms", "inline": True}, # [추가] 처리 시간 기록
                {"name": "공고 요약", "value": short_duties, "inline": False},
                {"name": "결과 요약", "value": short_result, "inline": False}
            ]
        }]
    }
    try: requests.post(webhook_url, json=payload)
    except: pass

# 별점 평가 기능
def send_discord_feedback(user_name, score, feedback_text):
    webhook_url = st.secrets.get("DISCORD_WEBHOOK_URL", None)
    if not webhook_url: return

    # 점수에 따른 색상 및 멘트 변화
    colors = {1: 15158332, 2: 15105570, 3: 16776960, 4: 3066993, 5: 5763719} # 빨주노초파
    comments = {1: "😭 최악이에요", 2: "😞 별로예요", 3: "😐 보통이에요", 4: "🙂 좋아요", 5: "😍 최고예요!"}

    payload = {
        "username": "⭐ 만족도 평가 알림",
        "embeds": [{
            "title": f"사용자 만족도 평가: {score}점 {'⭐' * score}",
            "description": f"**평가:** {comments.get(score, '')}",
            "color": colors.get(score, 3447003), # 점수별 색상
            "fields": [
                {"name": "사용자", "value": user_name if user_name else "익명", "inline": True},
                {"name": "추가 의견", "value": feedback_text if feedback_text else "없음", "inline": False}
            ]
        }]
    }
    try: requests.post(webhook_url, json=payload)
    except: pass

# ==========================================
# 이력서 읽기 & AI 분석
# ==========================================
def read_resume_file(uploaded_file):
    try:
        doc = docx.Document(uploaded_file)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip(): full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text.strip()]
                if row_text: full_text.append(" | ".join(row_text))
        return "\n".join(full_text)
    except Exception as e: return f"읽기 실패, 이력서 예시를 참고해주세요: {e}"

class AgentState(TypedDict):
    resume_text: str; duties: str; requirements: str; preferred: str; final_result: str

def match_node(state: AgentState):
    prompt = ChatPromptTemplate.from_template("""
    당신은 경력 20년차 HR 전문가입니다. 취업을 위한 경험 분석을 도와주세요. 
    [이력서] {resume}
    [공고] 업무: {duties}, 자격: {requirements}, 우대: {preferred}
    
    분석 결과 리포트:
    1. 적합도 점수 (0~100)
    2. 합격 포인트
    3. 보완 필요 사항
    4. 추천 미니 프로젝트
    """)
    chain = prompt | llm
    res = chain.invoke({
        "resume": state['resume_text'], "duties": state['duties'], 
        "requirements": state['requirements'], "preferred": state['preferred']
    })
    return {"final_result": res.content}

workflow = StateGraph(AgentState)
workflow.add_node("matcher", match_node)
workflow.set_entry_point("matcher")
workflow.add_edge("matcher", END)
app = workflow.compile()

# ==========================================
# 4. 화면(UI)
# ==========================================
st.title("🦦 AI 취업 도우미: 채용공고 이력서 매칭")

# [Session State 초기화] 분석 결과와 레이턴시 저장 공간 만들기
if 'analysis_result' not in st.session_state:
    st.session_state.analysis_result = None
if 'analysis_latency' not in st.session_state:
    st.session_state.analysis_latency = 0

with st.sidebar:
    st.header("사용자 정보")
    user_name = st.text_input("이름/닉네임", placeholder="예: 냐냥")

col1, col2 = st.columns([1, 1])

with col1:
    st.subheader("1. 내 이력서 : DOCX 파일")
    uploaded_file = st.file_uploader("Upload .docx", type=["docx"])

    st.markdown("---")
    st.subheader("2. 채용공고")
    duties_input = st.text_area("주요 업무", height=100)
    req_input = st.text_area("자격 요건", height=100)
    pref_input = st.text_area("우대 사항", height=100)

    run_btn = st.button("!! 분석 시작 !!", type="primary", use_container_width=True)

with col2:
    st.subheader("3. 분석 결과")
    
    # [분석] - 버튼 누르면 실행
    if run_btn:
        if not uploaded_file or (not duties_input and not req_input):
            st.warning("내용을 입력해주세요.")
        else:
            with st.spinner("AI가 채용공고를 분석하는 중이에요🐈‍⬛..."):
                resume_text = read_resume_file(uploaded_file)
                
                # [수정] AI Agent 호출 직전에 시간 측정 시작
                start_time = time.time()
                
                result = app.invoke({
                    "resume_text": resume_text, "duties": duties_input,
                    "requirements": req_input, "preferred": pref_input
                })
                
                # AI Agent 호출 직후 시간 측정 종료 및 계산 (밀리초 단위)
                end_time = time.time()
                latency_ms = round((end_time - start_time) * 1000)
                
                # 리셋 방지
                st.session_state.analysis_result = result['final_result']
                st.session_state.analysis_latency = latency_ms 
                
                # 분석 완료 - 디스코드 전송
                send_discord_alert(user_name, duties_input, result['final_result'], latency_ms)

    # [결과 출력] 저장된 결과가 있으면 보여줌
    if st.session_state.analysis_result:
        # [추가] 분석 시간을 UI에 표시
        st.caption(f"⏱️ 분석 시간: **{st.session_state.analysis_latency:,} ms**") 
        st.markdown(st.session_state.analysis_result)
        st.markdown("---")
        
        # ==========================================
        # 만족도 평가
        # ==========================================
        st.subheader("⭐ 분석 결과가 만족스러우신가요? 별점을 매겨주세요. 개발자에게 힘이 됩니다")
        
        # 별점 위젯 (stars)
        sentiment_mapping = ["1점", "2점", "3점", "4점", "5점"]
        selected = st.feedback("stars")
        
        if selected is not None:
            # selected는 0~4로 들어오므로 +1 해줌
            score = selected + 1
            st.toast(f"소중한 의견 감사합니다! ({score}점)", icon="🐈‍⬛")
            
            # 디스코드로 별점 전송
            # (중복 전송 방지 - 세션 키 확인)
            if 'feedback_sent' not in st.session_state or st.session_state.feedback_sent != score:
                send_discord_feedback(user_name, score, "사용자가 별점을 클릭했습니다.")
                st.session_state.feedback_sent = score
